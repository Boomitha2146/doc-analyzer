import os
import io
import json
import logging
from pathlib import Path
from typing import Optional

import anthropic
import pytesseract
import pdfplumber
from PIL import Image
from docx import Document
from fastapi import FastAPI, File, UploadFile, HTTPException, Security, Depends
from fastapi.security.api_key import APIKeyHeader
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="AI Document Analysis API",
    description="Extract, analyse, and summarise content from PDF, DOCX, and image files using AI.",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Auth ──────────────────────────────────────────────────────────────────────
API_KEY = os.environ.get("DOC_API_KEY", "hackathon-key-2024")
api_key_header = APIKeyHeader(name="X-API-Key", auto_error=False)


def verify_api_key(key: Optional[str] = Security(api_key_header)):
    if key != API_KEY:
        raise HTTPException(status_code=403, detail="Invalid or missing API key")
    return key


# ── Response model ────────────────────────────────────────────────────────────
class AnalysisResponse(BaseModel):
    filename: str
    file_type: str
    extracted_text_preview: str
    summary: str
    entities: dict
    sentiment: dict


# ── Text extraction helpers ───────────────────────────────────────────────────

def extract_from_pdf(data: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n\n".join(text_parts).strip()


def extract_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n\n".join(paragraphs).strip()


def extract_from_image(data: bytes) -> str:
    image = Image.open(io.BytesIO(data))
    # Upscale small images for better OCR
    w, h = image.size
    if w < 1000:
        scale = 1000 / w
        image = image.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    text = pytesseract.image_to_string(image, config="--psm 6")
    return text.strip()


def detect_file_type(filename: str, content_type: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" or content_type == "application/pdf":
        return "pdf"
    if ext in (".docx", ".doc") or "word" in content_type:
        return "docx"
    if ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp") or content_type.startswith("image/"):
        return "image"
    raise HTTPException(status_code=415, detail=f"Unsupported file type: {ext}")


# ── AI analysis ───────────────────────────────────────────────────────────────

ANALYSIS_PROMPT = """You are an expert document analyst. Analyse the following extracted document text and return a structured JSON response.

Text to analyse:
\"\"\"
{text}
\"\"\"

Return ONLY valid JSON (no markdown, no extra text) with exactly this structure:
{{
  "summary": "A concise 2–4 sentence summary of the document's main content and purpose.",
  "entities": {{
    "persons": ["list of person names found"],
    "organisations": ["list of organisation/company names found"],
    "locations": ["list of locations/places found"],
    "dates": ["list of dates or time references found"],
    "monetary_amounts": ["list of monetary values found"],
    "other": ["any other notable named entities"]
  }},
  "sentiment": {{
    "label": "positive | negative | neutral",
    "score": 0.0,
    "explanation": "One sentence explaining the sentiment classification."
  }}
}}

Rules:
- sentiment.score must be between -1.0 (very negative) and 1.0 (very positive), 0.0 for neutral.
- If the document is purely factual/informational with no emotional tone, classify as neutral (score near 0).
- Extract real entities only; do not hallucinate names not present in the text.
- Return empty arrays [] if no entities of that type are found.
"""


def analyse_with_ai(text: str) -> dict:
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    # Truncate to ~12 000 chars to stay well within context
    truncated = text[:12000]

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=1500,
        messages=[
            {"role": "user", "content": ANALYSIS_PROMPT.format(text=truncated)}
        ],
    )

    raw = message.content[0].text.strip()
    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw)


# ── Endpoint ──────────────────────────────────────────────────────────────────

@app.post("/analyze", response_model=AnalysisResponse)
async def analyze_document(
    file: UploadFile = File(...),
    _: str = Depends(verify_api_key),
):
    """
    Upload a PDF, DOCX, or image file and receive:
    - Extracted text preview
    - AI-generated summary
    - Named entity extraction (persons, orgs, locations, dates, monetary amounts)
    - Sentiment analysis with score and explanation
    """
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")

    file_type = detect_file_type(file.filename or "upload", file.content_type or "")
    logger.info(f"Processing {file.filename} as {file_type} ({len(data)} bytes)")

    # 1. Extract text
    try:
        if file_type == "pdf":
            text = extract_from_pdf(data)
        elif file_type == "docx":
            text = extract_from_docx(data)
        else:
            text = extract_from_image(data)
    except Exception as e:
        logger.error(f"Extraction error: {e}")
        raise HTTPException(status_code=422, detail=f"Text extraction failed: {str(e)}")

    if not text:
        raise HTTPException(status_code=422, detail="No readable text found in the document.")

    # 2. AI analysis
    try:
        analysis = analyse_with_ai(text)
    except json.JSONDecodeError as e:
        logger.error(f"AI JSON parse error: {e}")
        raise HTTPException(status_code=502, detail="AI returned malformed response.")
    except Exception as e:
        logger.error(f"AI analysis error: {e}")
        raise HTTPException(status_code=502, detail=f"AI analysis failed: {str(e)}")

    return AnalysisResponse(
        filename=file.filename or "unknown",
        file_type=file_type,
        extracted_text_preview=text[:500] + ("…" if len(text) > 500 else ""),
        summary=analysis["summary"],
        entities=analysis["entities"],
        sentiment=analysis["sentiment"],
    )


@app.get("/health")
def health():
    return {"status": "ok", "version": "1.0.0"}


@app.get("/")
def root():
    return {
        "service": "AI Document Analysis API",
        "docs": "/docs",
        "analyze_endpoint": "POST /analyze",
        "auth": "Pass your API key in the X-API-Key header",
    }
